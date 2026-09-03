from __future__ import annotations

import csv
import hashlib
import re
import shutil
from html.parser import HTMLParser
from io import StringIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".html", ".htm", ".docx", ".csv", ".xlsx"}
CURRENT_SUMMARY_EXTENSIONS = {".pdf", ".md", ".txt", ".html", ".htm"}
TRANSCRIPT_EXTENSIONS = {".txt", ".md"}
MARKDOWN_FOLDER_NAME = "md"


class SourceFileError(ValueError):
    pass


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in {"script", "style", "svg", "img", "picture", "video", "audio"}:
            self._ignored_depth += 1
        elif not self._ignored_depth and tag in {"p", "br", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "svg", "img", "picture", "video", "audio"}:
            self._ignored_depth = max(0, self._ignored_depth - 1)
        elif not self._ignored_depth and tag in {"p", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth:
            self.parts.append(data)

    def text(self) -> str:
        return "\n".join(line.strip() for line in "".join(self.parts).splitlines() if line.strip())


def safe_filename(name: str, fallback: str = "document") -> str:
    name = Path(name).name
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name).strip().rstrip(".")
    return cleaned or fallback


def summed_output_stem(source_name: str) -> str:
    stem = Path(source_name).stem
    if stem.casefold().endswith(".summary.raw"):
        stem = stem[: -len(".summary.raw")]
    if "요약본" in stem:
        before, after = stem.rsplit("요약본", 1)
        stem = f"{before}summed{after}"
    else:
        stem = f"{stem} summed"
    return safe_filename(stem, fallback="summed")


def preferred_summary_text_source(source: Path) -> Path:
    if source.suffix.casefold() != ".pdf":
        return source
    stem = source.stem
    base = re.sub(r"\s*요약본$", "", stem)
    companion = source.with_name(f"{base}.summary.raw.md")
    return companion if companion.is_file() else source


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def save_bytes(data: bytes, filename: str, folder: Path) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    destination = folder / safe_filename(filename)
    number = 2
    while destination.exists():
        destination = folder / f"{Path(filename).stem} ({number}){Path(filename).suffix}"
        number += 1
    destination.write_bytes(data)
    return destination


def copy_unique(source: Path, folder: Path) -> Path:
    folder.mkdir(parents=True, exist_ok=True)
    destination = folder / safe_filename(source.name)
    number = 2
    while destination.exists() and sha256_file(destination) != sha256_file(source):
        destination = folder / f"{source.stem} ({number}){source.suffix}"
        number += 1
    if not destination.exists():
        shutil.copy2(source, destination)
    return destination


def _extract_pdf(path: Path) -> str:
    if path.read_bytes()[:5] != b"%PDF-":
        raise SourceFileError(f"올바른 PDF가 아닙니다: {path.name}")
    reader = PdfReader(str(path))
    pages = []
    for number, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        pages.append(f"\n===== PDF {number}쪽 =====\n{text}")
    return "".join(pages)


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    rows = csv.reader(StringIO(text))
    return "\n".join("\t".join(cell.strip() for cell in row) for row in rows)


def _extract_xlsx(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"\n===== 시트: {sheet.title} =====")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append("\t".join(values))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    if not path.is_file():
        raise SourceFileError(f"파일을 찾을 수 없습니다: {path}")
    suffix = path.suffix.casefold()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise SourceFileError(f"지원하지 않는 형식입니다: {path.name}")
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix == ".csv":
        text = _extract_csv(path)
    elif suffix == ".xlsx":
        text = _extract_xlsx(path)
    elif suffix in {".html", ".htm"}:
        parser = _TextHTMLParser()
        parser.feed(path.read_text(encoding="utf-8", errors="replace"))
        text = parser.text()
    else:
        text = path.read_text(encoding="utf-8-sig", errors="replace")
    text = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    if len(re.sub(r"\s+", "", text)) < 80:
        raise SourceFileError(
            f"텍스트를 충분히 추출하지 못했습니다: {path.name}. "
            "스캔 이미지만 있는 파일은 이미지 제외 설정에서 처리할 수 없습니다."
        )
    return text


def write_extracted_text(source: Path, destination: Path) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(extract_text(source), encoding="utf-8")
    return destination


def meaningful_length(text: str) -> int:
    text = re.sub(r"=====.*?=====", "", text)
    return len(re.sub(r"\s+", "", text))
